import requests

import os
from io import BytesIO


def extract_text(file) -> str:
    ext = os.path.splitext(file.name)[-1].lower()

    file.seek(0)
    file_bytes = file.read()
    file_stream = BytesIO(file_bytes)

    if ext == ".pdf":
        return _extract_pdf(file_stream)
    elif ext == ".docx":
        return _extract_docx(file_stream)
    elif ext == ".xlsx":
        return _extract_xlsx(file_stream)
    elif ext in (".txt", ".csv"):
        return _extract_txt(file_bytes)
    else:
        raise ValueError(f"Qo'llab-quvvatlanmaydigan format: {ext}")


def _extract_pdf(stream):
    import fitz

    doc = fitz.open(stream=stream.read(), filetype="pdf")
    return "\n".join(page.get_text() for page in doc)


def _extract_docx(stream):
    from docx import Document
    doc = Document(stream)
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(paragraphs)

def _extract_xlsx(stream):
    from openpyxl import load_workbook
    wb = load_workbook(stream, data_only=True)
    lines = []
    for sheet in wb.worksheets:
        lines.append(f"=== {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            lines.append(" | ".join(str(cell) if cell else "" for cell in row))

    return "\n".join(lines)


def _extract_txt(file_bytes):
    return file_bytes.decode("utf-8", errors="ignore")


def check_file(file = None, text = None):
    try:
        txt = ""
        if file is not None:
            txt = extract_text(file)
        elif text is not None:
            txt = text
        data = {
            "txt_inp": txt
        }
        response = requests.post("https://plagiat.ai/check_txt2", data=data)
        return response.json(), True
    except requests.exceptions.RequestException:
        return "", False
